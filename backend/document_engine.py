"""
DocAI - Document Intelligence Engine
RAG pipeline: PDF/DOCX/TXT ingestion → ChromaDB → Groq LLM
"""

import os
import re
import json
import uuid
import hashlib
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, asdict

# Document parsers
import fitz  # PyMuPDF
import docx as python_docx
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Embeddings & Vector DB
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

# LLM
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate


# ─────────────────────────────────────────
#  Data Classes
# ─────────────────────────────────────────

@dataclass
class DocumentInfo:
    doc_id: str
    filename: str
    file_type: str
    page_count: int
    word_count: int
    chunk_count: int
    summary: str = ""

@dataclass
class QAResult:
    answer: str
    sources: List[Dict]
    confidence: str
    doc_ids_used: List[str]

@dataclass
class ExtractionResult:
    extracted_data: Dict
    raw_text: str
    doc_id: str

@dataclass
class ComparisonResult:
    similarities: List[str]
    differences: List[str]
    summary: str
    doc_a_name: str
    doc_b_name: str


# ─────────────────────────────────────────
#  Document Parser
# ─────────────────────────────────────────

class DocumentParser:
    """Parse PDF, DOCX, TXT files into clean text with metadata."""

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> Tuple[str, int]:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page_num, page in enumerate(doc, 1):
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"[Page {page_num}]\n{text}")
        full_text = "\n\n".join(text_parts)
        return full_text, len(doc)

    @staticmethod
    def parse_docx(file_bytes: bytes) -> Tuple[str, int]:
        import io
        doc = python_docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n[Tables]\n" + "\n".join(tables_text)
        page_estimate = max(1, len(full_text) // 3000)
        return full_text, page_estimate

    @staticmethod
    def parse_txt(file_bytes: bytes) -> Tuple[str, int]:
        text = file_bytes.decode("utf-8", errors="replace")
        page_estimate = max(1, len(text) // 3000)
        return text, page_estimate

    @classmethod
    def parse(cls, file_bytes: bytes, filename: str) -> Tuple[str, int]:
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return cls.parse_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            return cls.parse_docx(file_bytes)
        elif ext in (".txt", ".md", ".csv"):
            return cls.parse_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {ext}")


# ─────────────────────────────────────────
#  Document Intelligence Engine
# ─────────────────────────────────────────

PERSIST_DIR = "./chroma_db"

class DocumentEngine:

    def __init__(self, groq_api_key: str, model: str = "llama-3.3-70b-versatile"):
        self.groq_api_key = groq_api_key
        self.model = model
        self.documents: Dict[str, DocumentInfo] = {}

        # Embeddings (free, local)
        print("Loading embeddings model...")
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )

        # ChromaDB vector store
        self.vectorstore = Chroma(
            collection_name="docai_documents",
            embedding_function=self.embeddings,
            persist_directory=PERSIST_DIR,
        )

        # LLM
        self.llm = ChatGroq(
            model=model,
            temperature=0.2,
            max_tokens=4096,
            api_key=groq_api_key,
        )

        # Text splitter
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

        print("✅ Document engine ready")

    # ── Ingest ──────────────────────────────

    def ingest_document(self, file_bytes: bytes, filename: str) -> DocumentInfo:
        """Parse, chunk, embed, and store a document."""
        # Parse
        text, page_count = DocumentParser.parse(file_bytes, filename)

        if len(text.strip()) < 50:
            raise ValueError("Document appears to be empty or unreadable.")

        # Generate stable doc_id from content hash
        doc_id = hashlib.md5(file_bytes).hexdigest()[:12]
        file_type = Path(filename).suffix.lower().lstrip(".")

        # Check if already ingested
        if doc_id in self.documents:
            return self.documents[doc_id]

        # Chunk
        chunks = self.splitter.split_text(text)
        word_count = len(text.split())

        # Add to ChromaDB with metadata
        ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
        metadatas = [
            {"doc_id": doc_id, "filename": filename, "chunk_index": i, "file_type": file_type}
            for i in range(len(chunks))
        ]
        self.vectorstore.add_texts(texts=chunks, metadatas=metadatas, ids=ids)

        # Generate summary
        summary = self._generate_summary(text[:4000], filename)

        info = DocumentInfo(
            doc_id=doc_id,
            filename=filename,
            file_type=file_type,
            page_count=page_count,
            word_count=word_count,
            chunk_count=len(chunks),
            summary=summary,
        )
        self.documents[doc_id] = info
        return info

    def _generate_summary(self, text_preview: str, filename: str) -> str:
        try:
            response = self.llm.invoke([
                SystemMessage(content="You are a document analyst. Write a 2-3 sentence summary of the document. Be specific and factual."),
                HumanMessage(content=f"Document: {filename}\n\nContent preview:\n{text_preview}\n\nProvide a concise summary:"),
            ])
            return response.content.strip()
        except Exception:
            return "Document ingested successfully."

    # ── Q&A ─────────────────────────────────

    def answer_question(self, question: str, doc_ids: Optional[List[str]] = None) -> QAResult:
        """RAG: retrieve relevant chunks, answer with citations."""
        # Build filter
        where_filter = None
        if doc_ids:
            if len(doc_ids) == 1:
                where_filter = {"doc_id": doc_ids[0]}
            else:
                where_filter = {"doc_id": {"$in": doc_ids}}

        # Retrieve top-k chunks — try with filter first, fallback without
        try:
            results = self.vectorstore.similarity_search_with_relevance_scores(
                question,
                k=8,
                filter=where_filter,
            )
        except Exception:
            results = self.vectorstore.similarity_search_with_relevance_scores(
                question, k=8
            )

        # If no results with filter, get all chunks for those docs directly
        if not results and doc_ids:
            try:
                all_chunks = self.vectorstore.get(where={"doc_id": {"$in": doc_ids}} if len(doc_ids) > 1 else {"doc_id": doc_ids[0]})
                if all_chunks and all_chunks.get("documents"):
                    docs_text = "\n\n".join(all_chunks["documents"][:8])
                    results = []
                    # Build fake results from direct fetch
                    for i, (txt, meta) in enumerate(zip(all_chunks["documents"][:8], all_chunks["metadatas"][:8])):
                        from langchain_core.documents import Document
                        results.append((Document(page_content=txt, metadata=meta), 0.75))
            except Exception:
                pass

        if not results:
            return QAResult(
                answer="No document chunks found. Please re-upload your document and try again.",
                sources=[],
                confidence="low",
                doc_ids_used=doc_ids or [],
            )

        # Build context — remove score threshold so all chunks are used
        context_parts = []
        sources = []
        seen_chunks = set()

        for doc, score in results:
            chunk_id = doc.metadata.get("chunk_index", 0)
            filename = doc.metadata.get("filename", "Unknown")
            key = f"{filename}_{chunk_id}"

            if key not in seen_chunks:
                seen_chunks.add(key)
                context_parts.append(
                    f"[Source: {filename}, Section {chunk_id + 1}]\n{doc.page_content}"
                )
                sources.append({
                    "filename": filename,
                    "doc_id": doc.metadata.get("doc_id", ""),
                    "section": chunk_id + 1,
                    "relevance": round(score * 100, 1),
                    "preview": doc.page_content[:200] + "...",
                })

        context = "\n\n---\n\n".join(context_parts)

        # Determine confidence
        avg_score = sum(s["relevance"] for s in sources) / max(len(sources), 1)
        confidence = "high" if avg_score > 70 else "medium" if avg_score > 40 else "low"

        # Generate answer
        system_prompt = """You are DocAI, an expert document analyst. Document content has been retrieved and is provided below. Use it to answer thoroughly.

Rules:
- ALWAYS answer from the document context — never say there is no context or no document provided
- For summaries: cover the main topic, all key points, important details and conclusions
- Be specific and mention which section info comes from
- Use bullet points for lists
- Format numbers and dates clearly
- Be detailed and complete — give a thorough answer"""

        response = self.llm.invoke([
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"Question: {question}\n\nDocument Content:\n{context}\n\nProvide a thorough answer based on the document content above:"),
        ])

        used_ids = list(set(s["doc_id"] for s in sources))
        return QAResult(
            answer=response.content.strip(),
            sources=sources[:4],
            confidence=confidence,
            doc_ids_used=used_ids,
        )

    # ── Structured Extraction ────────────────

    def extract_structured(self, doc_id: str, extraction_type: str) -> ExtractionResult:
        """Extract structured data: invoice fields, key dates, contact info, etc."""
        if doc_id not in self.documents:
            raise ValueError(f"Document {doc_id} not found.")

        # Get all chunks for this doc
        results = self.vectorstore.similarity_search(
            extraction_type, k=10, filter={"doc_id": doc_id}
        )
        full_text = "\n\n".join(r.page_content for r in results)

        extraction_prompts = {
            "invoice": """Extract invoice data as JSON with these fields:
{"invoice_number": "", "date": "", "due_date": "", "vendor_name": "", "vendor_address": "",
"client_name": "", "client_address": "", "line_items": [{"description": "", "quantity": "", "unit_price": "", "total": ""}],
"subtotal": "", "tax": "", "total_amount": "", "payment_terms": "", "currency": ""}""",

            "contract": """Extract contract key terms as JSON:
{"parties": [], "effective_date": "", "expiry_date": "", "contract_value": "",
"payment_terms": "", "key_obligations": [], "termination_clauses": [], "governing_law": "",
"notice_period": "", "penalties": [], "renewal_terms": ""}""",

            "resume": """Extract resume information as JSON:
{"full_name": "", "email": "", "phone": "", "location": "", "linkedin": "",
"summary": "", "skills": [], "experience": [{"company": "", "role": "", "duration": "", "highlights": []}],
"education": [{"institution": "", "degree": "", "year": ""}], "certifications": []}""",

            "general": """Extract key information as JSON with fields that make sense for this document type.
Include: document_type, main_topic, key_entities, important_dates, key_numbers, action_items, summary""",
        }

        prompt = extraction_prompts.get(extraction_type, extraction_prompts["general"])

        response = self.llm.invoke([
            SystemMessage(content=f"You are a data extraction specialist. Extract structured data from documents. Return ONLY valid JSON, no explanation.\n\n{prompt}"),
            HumanMessage(content=f"Document text:\n{full_text[:6000]}\n\nExtract and return JSON:"),
        ])

        raw = response.content.strip()
        # Clean JSON
        if "```json" in raw:
            raw = raw.split("```json")[1].split("```")[0].strip()
        elif "```" in raw:
            raw = raw.split("```")[1].split("```")[0].strip()

        try:
            extracted = json.loads(raw)
        except json.JSONDecodeError:
            extracted = {"raw_extraction": raw, "note": "Could not parse as structured JSON"}

        return ExtractionResult(
            extracted_data=extracted,
            raw_text=full_text[:500],
            doc_id=doc_id,
        )

    # ── Comparison ───────────────────────────

    def compare_documents(self, doc_id_a: str, doc_id_b: str) -> ComparisonResult:
        """Compare two documents and highlight similarities and differences."""
        doc_a = self.documents.get(doc_id_a)
        doc_b = self.documents.get(doc_id_b)

        if not doc_a or not doc_b:
            raise ValueError("One or both documents not found.")

        # Get representative chunks from each
        chunks_a = self.vectorstore.similarity_search(
            "main content key points terms", k=6, filter={"doc_id": doc_id_a}
        )
        chunks_b = self.vectorstore.similarity_search(
            "main content key points terms", k=6, filter={"doc_id": doc_id_b}
        )

        text_a = "\n\n".join(r.page_content for r in chunks_a)
        text_b = "\n\n".join(r.page_content for r in chunks_b)

        response = self.llm.invoke([
            SystemMessage(content="""You are a document comparison expert. Compare two documents and return a JSON with:
{"similarities": ["list of key similarities"], "differences": ["list of key differences"], "summary": "2-3 sentence overall comparison"}
Return ONLY valid JSON."""),
            HumanMessage(content=f"""Document A ({doc_a.filename}):\n{text_a[:3000]}\n\n---\n\nDocument B ({doc_b.filename}):\n{text_b[:3000]}\n\nCompare these documents:"""),
        ])

        raw = response.content.strip()
        if "```json" in raw:
            raw = raw.split("```json")[1].split("```")[0].strip()
        elif "```" in raw:
            raw = raw.split("```")[1].split("```")[0].strip()

        try:
            data = json.loads(raw)
        except Exception:
            data = {"similarities": [], "differences": [raw], "summary": raw}

        return ComparisonResult(
            similarities=data.get("similarities", []),
            differences=data.get("differences", []),
            summary=data.get("summary", ""),
            doc_a_name=doc_a.filename,
            doc_b_name=doc_b.filename,
        )

    # ── Utils ────────────────────────────────

    def remove_document(self, doc_id: str):
        """Remove a document and all its chunks from the vector store."""
        results = self.vectorstore.get(where={"doc_id": doc_id})
        if results and results.get("ids"):
            self.vectorstore.delete(ids=results["ids"])
        self.documents.pop(doc_id, None)

    def get_document_info(self, doc_id: str) -> Optional[DocumentInfo]:
        return self.documents.get(doc_id)

    def list_documents(self) -> List[DocumentInfo]:
        return list(self.documents.values())